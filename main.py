import streamlit as st
import time
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

#load variables from .env file
from dotenv import load_dotenv
import os
load_dotenv()

st.set_page_config(page_title="Unstop Drive Genius", page_icon="💻")

#get variables from .env file
account_id = os.getenv("ACCOUNT_ID")
API_TOKEN = "-m9pU-NT4cSsvJFsL3Q4Hg1imSpFclvtCKf8ANbM"
print(API_TOKEN)
def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a software engineer who generates the driver code in Java based on the given input and output format. A driver code has the code to take input from the user based on the given input format and then calls another function with the parameters as the input received and then display the output in the given output format. The code only contains the function description and not the logic, the code inside the function. Display the output as the value returned from the function in the given output format. " },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/mistral/mistral-7b-instruct-v0.1", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Driver Code", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("UnstopDriveGenius.docx")
    

def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "Unstop Drive Genius"
  
# left_co, cent_co,last_co = st.columns(3)
# with cent_co:
#     st.image("images/logo.png", width=200)

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    
st.text(
    "Fill in the input and output format to generate the driver code for the given input and output format."
)

st.subheader("Enter the input format")
input_format = st.text_area("Input format")

st.subheader("Enter the output format")
output_format = st.text_area("Output format")

if input_format:
    if output_format:
        if st.button("Generate Driver Code"):
            st.success("Driver code generated successfully", icon='📃')
            if "conversation" not in st.session_state:
                st.session_state.conversation = []
            user_input = f"Input format: {input_format}\nOutput format: {output_format}"
            doc = docx.Document()
            doc.save("UnstopDriveGenius.docx")
            bot_reply = bot_response(user_input)
            st.subheader("Driver Code")
            st.write(bot_reply)
            st.session_state.conversation.append((user_input, bot_reply))
        
            save_to_doc(st.session_state.conversation)
    
            st.download_button(
                label="Download driver code!",
                data=open("UnstopDriveGenius.docx", "rb").read(),
                file_name="Driver_Code.docx",
                mime="application/octet-stream",
                help="Click to download the driver code."
            )
    else:
        st.warning("Please enter the output format.")
else:
    st.warning("Please enter the input format.")
    